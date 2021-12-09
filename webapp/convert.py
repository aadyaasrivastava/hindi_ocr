import easyocr
from numpy import asarray
from PIL import Image
import docx
import uuid
from docx.shared import Pt
from fpdf import FPDF, fpdf
DOWNLOAD_FOLDER = '../webapp/download/document/'
def converter(img):
    image = Image.open(img)
    data = asarray(image)
    reader = easyocr.Reader(['hi'])
    result = reader.readtext(data,detail = 1)
    ans = ""
    for i in result:
        ans+=i[1]
        ans+=' '
    doc = docx.Document()
    para = doc.add_paragraph().add_run(ans)
    para.font.size = Pt(12)
    filename = str(uuid.uuid4())
    doc.save(DOWNLOAD_FOLDER+ "filename"+".docx")
DOWNLOAD_FOLDER = '../webapp/download/pdf/'
def converter_1(img):
    image = Image.open(img)
    data = asarray(image)
    reader = easyocr.Reader(['hi'])
    result = reader.readtext(data,detail = 1)
    ans = ""
    for i in result:
        ans+=i[1]
        ans+=' '
    document=FPDF()
    document.add_page()
    document.cell(400, 50, ln=1, align="L")
    document.output("pdf_file_name.pdf")
    document=FPDF(orientation='P', unit='mm', format='A3''A2''A1')
    fpdf.save(DOWNLOAD_FOLDER+ "filename"+".pdf")